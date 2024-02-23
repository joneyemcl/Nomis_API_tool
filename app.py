"""
Runner script
"""

import pandas as pd
from docx import Document


from definitions import DocStyles, PathsFile
from impact_areas import gen_impact_areas, ImpactAreas
from logger_setup import setup_logger
from menu import get_user_impact_area, get_user_receptors



# pandas display settings
pd.set_option('display.max_columns', 500)
pd.set_option('display.width', 500)

# setting up the logger
logger = setup_logger(__name__, reset_file=True)

# setting up doc
doc = Document(PathsFile.WORD_TEMPLATE)
para_title = doc.paragraphs[0]
para_title.style = DocStyles.heading2
run_para_title = para_title.add_run('Baseline')

# setting up excel workbook
writer = pd.ExcelWriter(PathsFile.EXCEL_OUTPUT, engine='xlsxwriter')

# getting user input (impact area)
impact_area = get_user_impact_area()

print(f'this is impact_area {impact_area}')
# creating impact area benchmark data structure
impact_areas = gen_impact_areas(impact_area)
impact_areas = ImpactAreas(impact_areas)
logger.debug(impact_areas)

# getting user receptors
user_receptors = get_user_receptors()

# executing the selected receptor modules
if user_receptors['myede'][1] == 'y':
    from processing.myede_processing import myede_processing
    from processing.myede_doc import add_myede_doc
    doc_inputs_myede = myede_processing(impact_areas, writer)
    doc = add_myede_doc(doc, doc_inputs_myede, impact_areas)

if user_receptors['snpp'][1] == 'y':
    from receptors.snpp.snpp_cleaning import get_cleaned_snpp
    from receptors.snpp.snpp_processing import snpp_processing
    from receptors.snpp.snpp_doc import add_snpp_to_doc
    snpp = get_cleaned_snpp()
    doc_inputs_snpp = snpp_processing(snpp, impact_areas, writer)
    add_snpp_to_doc(doc, doc_inputs_snpp, impact_areas.impact_area)

if user_receptors['bres'][1] == 'y':
    from receptors.bres.bres_cleaning import get_cleaned_bres
    from receptors.bres.bres_processing import bres_processing
    from receptors.bres.bres_doc import add_bres_to_doc
    bres = get_cleaned_bres()
    doc_inputs_bres = bres_processing(bres, impact_areas, writer)
    add_bres_to_doc(doc, doc_inputs_bres)

if user_receptors['unempl'][1] == 'y':
    from receptors.unemployment.unempl_cleaning import get_cleaned_unempl
    from receptors.unemployment.unempl_processing import unempl_processing
    from receptors.unemployment.unempl_doc import add_unempl_to_doc
    unempl = get_cleaned_unempl()
    doc_inputs_unempl = unempl_processing(unempl, impact_areas, writer)
    add_unempl_to_doc(doc, doc_inputs_unempl)

if user_receptors['occ'][1] == 'y':
    from receptors.occupation.occ_cleaning import get_cleaned_occ
    from receptors.occupation.occ_processing import occ_processing
    from receptors.occupation.occ_doc import add_occ_to_doc
    occ = get_cleaned_occ()
    doc_inputs_occ = occ_processing(occ, impact_areas, writer)
    add_occ_to_doc(doc, doc_inputs_occ)

if user_receptors['skills'][1] == 'y':
    from receptors.skills.skills_cleaning import get_cleaned_skills
    from receptors.skills.skills_processing import skills_processing
    from receptors.skills.skills_doc import add_skills_to_doc
    skills = get_cleaned_skills()
    doc_inputs_skills = skills_processing(skills, impact_areas, writer)
    add_skills_to_doc(doc, doc_inputs_skills)

if user_receptors['afford'][1] == 'y':
    from receptors.affordability.affordability_cleaning import get_cleaned_afford
    from receptors.affordability.affordability_processing import afford_processing
    from receptors.affordability.affordability_doc import add_aff_to_doc
    afford = get_cleaned_afford()
    doc_inputs_aff = afford_processing(afford, impact_areas, writer)
    add_aff_to_doc(doc, doc_inputs_aff, ImpactAreas.impact_area)

doc.save(PathsFile.WORD_OUTPUT)
writer._save()
writer.close()
print('\nOutput processed!')