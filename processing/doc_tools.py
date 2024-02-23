"""
Provides a selection of tools which take python data structures and objects and add them to a word document using the
python-docx module
"""

import logging
from typing import List

from docx import Document, shared

from definitions import DocStyles

logger = logging.getLogger(__name__)


def add_table_from_list(
        doc: Document, title: str, columns: List[str], rows: List[str], data: List[List], footnote: str) -> Document:
    """
    add a table with row and column headers to a word document and return the document
    data argument format:
    [row, row, row]
    """
    logger.info('Adding table from list to doc')
    # adding title
    doc.add_paragraph(text=title, style=DocStyles.tableTitle)

    # adding table
    no_rows = len(data)
    no_cols = len(columns)
    assert no_cols == len(data[0]), 'The number of headers and columns of data are not equal'
    assert no_rows == len(data), 'The number of row headers and rows of data are not equal'

    table = doc.add_table(no_rows + 1, no_cols + 1, style=DocStyles.table)

    # populating column headers
    for j in range(no_cols):
        table.cell(0, j+1).text = columns[j]
        table.cell(0, j+1).paragraphs[0].style = DocStyles.tableHeading

    # populating row headers
    for i in range(no_rows):
        table.cell(i+1, 0).text = rows[i]
        table.cell(i+1, 0).paragraphs[0].style = DocStyles.tableText

    # populating data cells
    for i in range(no_rows):
        for j in range(no_cols):
            table.cell(i+1, j+1).text = str(data[i][j])
            table.cell(i+1, j+1).paragraphs[0].style = DocStyles.tableNumber

    # adding a footnote
    doc.add_paragraph(text=footnote, style=DocStyles.tableSource)
    logger.info('Finished adding table from list to doc')

    return doc


def add_figure(doc: Document, title: str, img_path: str, source: str) -> Document:
    # adding a figure table
    fig_table = doc.add_table(3, 1, style=DocStyles.placeholder)

    # figure title
    fig_table.cell(0, 0).text = title
    fig_table.cell(0, 0).paragraphs[0].style = DocStyles.figureTitle

    # figure image
    para = fig_table.cell(1, 0).paragraphs[0]
    para.style = DocStyles.normalFigures
    run = para.add_run()
    run.add_picture(img_path, width=shared.Cm(15.2))

    # figure source
    fig_table.cell(2, 0).text = source
    fig_table.cell(2, 0).paragraphs[0].style = DocStyles.figureSource

    return doc
