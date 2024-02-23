"""
takes in processed myede text/data and adds it to a python-docx document
"""

import logging
from typing import Dict

from docx import Document

from definitions import DocStyles, AgeBands, FootNotes
from impact_areas import ImpactAreas
from processing.doc_tools import add_table_from_list
from processing.myede_processing import year_recent, year_less_ten

logger = logging.getLogger(__name__)


# populating the word document
def add_myede_doc(doc: Document, doc_inputs: Dict, impact_areas: ImpactAreas) -> Document:
    """
    Add myede content to the word document and return the document
    """
    logger.info("Adding myede content to doc")
    # add title
    doc.add_paragraph(text='Population', style=DocStyles.heading4)
    # paragraph one
    doc.add_paragraph(text=doc_inputs['text_0'], style=DocStyles.heading5)
    # paragraph two
    doc.add_paragraph(text=doc_inputs['text_1'], style=DocStyles.heading5)
    # add table
    kwargs = {
        'title': f'Change in population by age, {year_less_ten}/{str(year_recent)[2:]}',
        'columns': impact_areas.area_names,
        'rows': AgeBands.bands,
        'data': doc_inputs['table_pct_chg'],
        'footnote': FootNotes.MYEDE,
    }
    doc = add_table_from_list(doc, **kwargs)
    logger.info("Finished adding myede content to doc")
    return doc
